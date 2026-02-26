"""
전북지역 졸업생 네트워크 - 사용자 매뉴얼 DOCX 생성
실행: python3 scripts/generate-manual.py
전제조건: manual-screenshots/ 폴더에 스크린샷이 있어야 함
"""

import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SCREENSHOT_DIR = os.path.join(BASE_DIR, 'manual-screenshots')
OUTPUT_FILE = os.path.join(BASE_DIR, '전북지역_졸업생네트워크_사용자매뉴얼.docx')

BRAND_BLUE  = RGBColor(0x1e, 0x40, 0xaf)   # #1e40af
BRAND_LIGHT = RGBColor(0xdb, 0xe9, 0xfe)   # #dbe9fe
GRAY        = RGBColor(0x6b, 0x72, 0x80)   # #6b7280
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)


def add_page_break(doc):
    doc.add_page_break()


def set_heading_style(paragraph, level, text, color=None):
    """제목 스타일 적용"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    run.font.bold = True
    if color:
        run.font.color.rgb = color
    if level == 0:
        run.font.size = Pt(22)
    elif level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)


def add_screenshot(doc, img_path, caption, width=Inches(5.5)):
    """스크린샷 이미지 + 캡션 삽입"""
    if img_path and os.path.exists(img_path):
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_path, width=width)
        except Exception as e:
            p = doc.add_paragraph(f'[이미지 로드 실패: {e}]')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f'[스크린샷 없음: {os.path.basename(img_path) if img_path else ""}]')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = GRAY

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.italic = True
    cap.runs[0].font.color.rgb = GRAY
    doc.add_paragraph()  # 여백


def add_info_box(doc, text, bg_type='info'):
    """안내 박스 텍스트 삽입"""
    prefix = '💡 ' if bg_type == 'info' else '⚠️ '
    p = doc.add_paragraph()
    run = p.add_run(f'{prefix}{text}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1e, 0x56, 0xa0) if bg_type == 'info' else RGBColor(0x92, 0x40, 0x0e)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_steps(doc, steps):
    """순서 있는 단계 목록"""
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f'{i}. {step}', style='List Number')
        p.paragraph_format.left_indent = Cm(0.5)
        p.runs[0].font.size = Pt(10.5)


def add_bullets(doc, items):
    """글머리 기호 목록"""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.5)
        p.runs[0].font.size = Pt(10.5)


def img(name):
    """스크린샷 파일 경로 반환"""
    path = os.path.join(SCREENSHOT_DIR, f'{name}.png')
    return path if os.path.exists(path) else None


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = BRAND_BLUE
    p.runs[0].font.size = Pt(16)
    p.paragraph_format.space_before = Pt(16)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    p.runs[0].font.size = Pt(13)
    p.paragraph_format.space_before = Pt(10)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10.5)
    return p


# ──────────────────────────────────────────────────────────────────
# 메인
# ──────────────────────────────────────────────────────────────────

def build_manual():
    doc = Document()

    # 페이지 여백 설정
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── 표지 ─────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('전북지역', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = BRAND_BLUE
    title.runs[0].font.size = Pt(24)

    sub = doc.add_heading('졸업생 네트워크 플랫폼', level=0)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = BRAND_BLUE
    sub.runs[0].font.size = Pt(22)

    desc = doc.add_heading('사 용 자 매 뉴 얼', level=0)
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.runs[0].font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    desc.runs[0].font.size = Pt(18)

    doc.add_paragraph()
    doc.add_paragraph()

    add_screenshot(doc, img('01_main'), '메인 화면', width=Inches(4.5))

    doc.add_paragraph()
    ver_p = doc.add_paragraph('버전 v2.0  |  2026년 2월 26일  |  IT팀')
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_p.runs[0].font.color.rgb = GRAY
    ver_p.runs[0].font.size = Pt(10)

    add_page_break(doc)

    # ── 목차 ─────────────────────────────────────────────────────
    h1(doc, '목   차')
    toc_items = [
        ('1장', '시작하기', '로그인 · 회원가입'),
        ('2장', '학생 기능', '대시보드 · 채용정보 · 취업박람회 · 산업체견학 · 자격증지원 · 진로상담 · 동문네트워킹 · 프로필 · 경력관리'),
        ('3장', '교사 기능', '대시보드 · 채용공고관리 · 지원자관리 · 진로상담관리'),
        ('4장', '관리자 기능', '회원관리 · 채용공고관리 · 게시판관리 · 공지사항관리 · 코드관리'),
    ]
    for num, title_text, detail in toc_items:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{num}  {title_text}')
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(f'\n      {detail}')
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = GRAY
        p.paragraph_format.space_after = Pt(6)

    add_page_break(doc)

    # ╔═══════════════════════════════════════════╗
    # ║   1장: 시작하기                           ║
    # ╚═══════════════════════════════════════════╝
    h1(doc, '1장. 시작하기')

    # 로그인
    h2(doc, '1.1 로그인')
    body(doc, '등록된 이메일과 비밀번호를 입력하여 서비스에 접속합니다.')
    add_screenshot(doc, img('02_login'), '▲ 로그인 화면')
    add_steps(doc, [
        '웹 브라우저에서 jjobb.kr 접속',
        '"로그인" 버튼 클릭',
        '이메일 · 비밀번호 입력',
        '"로그인" 버튼 클릭 → 대시보드 이동',
    ])
    add_info_box(doc, 'Mac(macOS) 사용자 주의: 이메일 입력 시 첫 글자가 자동으로 대문자로 변환될 수 있습니다. 이메일은 반드시 소문자로 입력해주세요.', 'warning')

    # 회원가입
    h2(doc, '1.2 회원가입')
    body(doc, '재학생·졸업생은 아래 절차로 계정을 생성합니다.')
    add_screenshot(doc, img('03_register'), '▲ 회원가입 화면')
    add_steps(doc, [
        '메인 페이지에서 "회원가입" 버튼 클릭',
        '이름, 이메일, 비밀번호(6자 이상), 전화번호, 학교, 사용자 유형, 졸업년도·전공 입력',
        '"가입하기" 클릭 → 서버에 저장 후 로그인 페이지 이동',
    ])
    add_info_box(doc, '가입 정보는 서버 DB에 저장되므로 어떤 기기에서도 동일 계정으로 로그인할 수 있습니다.')

    add_page_break(doc)

    # ╔═══════════════════════════════════════════╗
    # ║   2장: 학생 기능                          ║
    # ╚═══════════════════════════════════════════╝
    h1(doc, '2장. 학생 기능')

    # 대시보드
    h2(doc, '2.1 대시보드')
    body(doc, '로그인 후 첫 화면입니다. 신규 채용공고, 지원 현황, 최근 소식, 교육 프로그램을 한눈에 확인합니다.')
    add_screenshot(doc, img('04_student_dashboard'), '▲ 학생 대시보드')

    # 채용정보
    h2(doc, '2.2 채용 정보')
    body(doc, '기업별 채용공고를 검색하고 지원할 수 있습니다.')
    add_screenshot(doc, img('05_student_jobs'), '▲ 채용 정보 목록')
    add_steps(doc, [
        '메뉴 → "채용 정보" 클릭',
        '키워드·지역·고용형태·경력 필터로 공고 검색',
        '원하는 공고 클릭 → 상세 정보 확인',
        '"지원하기" 버튼 클릭 → 지원 완료',
    ])

    # 취업박람회
    h2(doc, '2.3 취업박람회')
    body(doc, '지역 기업이 참여하는 취업박람회 정보를 확인하고 참가 신청을 합니다.')
    add_screenshot(doc, img('06_student_jobfair'), '▲ 취업박람회 목록')
    add_steps(doc, [
        '메뉴 → "취업박람회" 클릭',
        '박람회 카드에서 일정·장소·참가비 확인',
        '"참가 신청" 버튼 클릭',
        '이름·연락처·이메일·메시지 입력 후 "신청 확인"',
    ])
    add_info_box(doc, '로그인 상태에서는 이름과 이메일이 자동으로 입력됩니다.')

    # 산업체 견학
    h2(doc, '2.4 산업체 견학')
    body(doc, '지역 기업 현장 견학 프로그램을 신청합니다.')
    add_screenshot(doc, img('07_student_industryvisit'), '▲ 산업체 견학 목록')
    add_steps(doc, [
        '메뉴 → "산업체 견학" 클릭',
        '견학 기업 카드에서 모집인원·혜택 확인',
        '"견학 신청" 버튼 클릭 후 정보 입력 → 신청 완료',
    ])

    # 자격증 지원
    h2(doc, '2.5 자격증 지원')
    body(doc, '자격증 취득 지원 프로그램(수강료 지원, 무료 강의 등)을 신청합니다.')
    add_screenshot(doc, img('08_student_certification'), '▲ 자격증 지원 목록')
    add_steps(doc, [
        '메뉴 → "자격증 지원" 클릭',
        '유형별 필터(기사/기능사/IT 등)로 검색',
        '"지원 신청" 버튼 클릭 후 신청 정보 입력',
    ])

    # 진로 상담
    h2(doc, '2.6 진로 상담')
    body(doc, '담당 교사에게 진로·취업 관련 상담을 신청하고 답변을 받습니다.')
    add_screenshot(doc, img('09_student_counseling'), '▲ 진로 상담 신청')
    add_steps(doc, [
        '메뉴 → "진로 상담" 클릭',
        '"상담 신청하기" 버튼 클릭',
        '상담 유형 선택(진로/취업/학업/기타), 제목·내용 입력',
        '"신청하기" 클릭',
        '교사 승인 후 "승인됨" 상태에서 답변 확인',
    ])
    doc.add_paragraph('상담 상태 안내:')
    add_bullets(doc, [
        '🟡 대기중 – 교사 처리 전',
        '🟢 승인됨 – 답변 작성 완료',
        '🔴 거절됨 – 거절 사유 확인',
    ])

    add_page_break(doc)

    # 동문 네트워킹
    h2(doc, '2.7 동문 네트워킹')
    body(doc, '재학생·졸업생과 네트워크를 형성하고 메시지를 주고받습니다.')
    add_screenshot(doc, img('10_student_networking'), '▲ 동문 네트워킹')
    add_steps(doc, [
        '메뉴 → "동문 네트워킹" 클릭',
        '이름·졸업년도·회사·직무로 동문 검색',
        '"네트워크 추가" 클릭 → 연결 요청',
        '연결된 동문에게 "메시지 보내기" 가능',
    ])

    # 내 프로필
    h2(doc, '2.8 내 프로필')
    body(doc, '전화번호, 자기소개, 관심 분야, SNS 링크 등 프로필 정보를 수정합니다.')
    add_screenshot(doc, img('11_student_profile'), '▲ 내 프로필')
    add_steps(doc, [
        '메뉴 → "내 프로필" → "프로필 수정"',
        '원하는 항목 수정 후 "저장"',
    ])

    # 경력 관리
    h2(doc, '2.9 경력 관리')
    body(doc, '경력·자격증·교육 이력·포트폴리오를 등록하여 채용 지원 시 자동으로 활용됩니다.')
    add_screenshot(doc, img('12_student_career'), '▲ 경력 관리')
    add_steps(doc, [
        '메뉴 → "경력 관리" 클릭',
        '탭 선택: 경력 / 자격증 / 교육 / 포트폴리오',
        '"추가" 버튼 클릭 → 정보 입력 → "저장"',
    ])

    add_page_break(doc)

    # ╔═══════════════════════════════════════════╗
    # ║   3장: 교사 기능                          ║
    # ╚═══════════════════════════════════════════╝
    h1(doc, '3장. 교사 기능')
    body(doc, '교사 계정으로 로그인하면 학생 상담 관리 및 채용공고 등록·관리 기능이 제공됩니다.')

    # 교사 대시보드
    h2(doc, '3.1 교사 대시보드')
    add_screenshot(doc, img('13_teacher_dashboard'), '▲ 교사 대시보드')
    add_bullets(doc, [
        '내 채용공고 목록 확인',
        '최근 지원자 목록 조회',
        '대기 중인 상담 요청 확인',
        '최근 소식 · 교육 프로그램 확인',
    ])

    # 채용공고 등록
    h2(doc, '3.2 채용공고 등록 및 관리')
    add_steps(doc, [
        '대시보드 → "공고 등록" 클릭',
        '회사명, 직무명, 근무지역, 급여, 고용형태, 경력/학력, 직무설명, 자격요건, 복리후생, 마감일 입력',
        '"등록하기" 클릭',
        '수정/삭제: 내 채용공고 목록에서 수정·삭제 버튼 사용',
    ])
    add_info_box(doc, '채용공고를 삭제하면 해당 공고의 지원자 정보도 함께 삭제됩니다.', 'warning')

    # 진로 상담 관리
    h2(doc, '3.3 진로 상담 관리')
    add_screenshot(doc, img('14_teacher_counseling'), '▲ 교사 진로 상담 관리')
    add_steps(doc, [
        '메뉴 → "진로 상담" 클릭',
        '대기 중인 상담 요청 선택',
        '"승인" (답변 작성) 또는 "거절" (사유 작성) 선택',
        '"저장" 클릭 → 학생에게 결과 전달',
    ])

    add_page_break(doc)

    # ╔═══════════════════════════════════════════╗
    # ║   4장: 관리자 기능                        ║
    # ╚═══════════════════════════════════════════╝
    h1(doc, '4장. 관리자 기능')
    body(doc, '관리자 계정으로 로그인하면 시스템 전체 데이터를 관리하는 메뉴가 제공됩니다.')

    # 관리자 대시보드
    h2(doc, '4.1 관리자 대시보드')
    add_screenshot(doc, img('15_admin_dashboard'), '▲ 관리자 대시보드')
    add_bullets(doc, [
        '전체 사용자 수 · 활성 채용공고 수 · 상담 건수 통계',
        '최근 소식 및 교육 프로그램 현황',
    ])

    # 회원 관리
    h2(doc, '4.2 회원 관리')
    body(doc, '전체 회원 정보를 조회, 수정, 삭제합니다.')
    add_screenshot(doc, img('16_admin_users'), '▲ 회원 관리')
    add_steps(doc, [
        '메뉴 → "회원 관리" 클릭',
        '유형별 필터(전체/학생/졸업생/교사/기업/관리자) 및 이름·이메일 검색',
        '"수정" 클릭 → 이름, 이메일, 전화번호, 유형, 학교, 졸업년도 수정 → "저장"',
        '"삭제" 클릭 → 확인 → 해당 회원 완전 삭제',
    ])
    add_info_box(doc, '회원을 삭제하면 해당 회원의 모든 데이터(상담·지원 정보)가 함께 삭제됩니다.', 'warning')

    # 채용공고 관리
    h2(doc, '4.3 채용공고 관리')
    body(doc, '전체 채용공고를 등록, 수정, 삭제, 만료 처리합니다.')
    add_screenshot(doc, img('17_admin_jobs'), '▲ 채용공고 관리')
    add_steps(doc, [
        '메뉴 → "채용공고 관리" 클릭',
        '상태별 필터: 전체 / 활성 / 마감 / 만료',
        '"공고 등록" 버튼으로 신규 공고 등록',
        '공고 선택 → "수정" 또는 "삭제"',
        '마감일 지난 공고 → "만료" 처리 후 학생 목록에서 제외',
    ])

    # 게시판 관리
    h2(doc, '4.4 게시판 관리 (최근 소식 · 교육 프로그램)')
    body(doc, '대시보드에 표시되는 최근 소식과 평생교육 프로그램을 관리합니다.')
    add_screenshot(doc, img('18_admin_board'), '▲ 게시판 관리')
    add_steps(doc, [
        '메뉴 → "게시판 관리" 클릭',
        '"최근 소식" 탭: "소식 등록" → 제목·내용 입력 → "등록"',
        '"평생교육프로그램" 탭: "프로그램 등록" → 프로그램명·설명·기간·수강료 입력 → "등록"',
        '목록에서 "수정" / "삭제" 버튼으로 관리',
    ])

    add_page_break(doc)

    # 공지사항 관리
    h2(doc, '4.5 공지사항 관리 (취업박람회 · 산업체견학 · 자격증지원)')
    body(doc, '학생들이 조회·신청하는 취업박람회, 산업체 견학, 자격증 지원 정보를 등록·관리합니다.')
    add_screenshot(doc, img('19_admin_announcements'), '▲ 공지사항 관리')
    add_steps(doc, [
        '메뉴 → "공지사항 관리" 클릭',
        '"새 공지사항 등록" 클릭',
        '유형 선택: 취업박람회 / 산업체견학 / 자격증지원',
        '제목, 주최기관, 행사일, 장소, 마감일, 모집인원, 참가비, 혜택, 참가조건, 연락처 입력',
        '(선택) 이미지 URL · 상세 URL 입력',
        '"저장" 클릭 → 학생 페이지에 즉시 반영',
    ])
    doc.add_paragraph('수정 / 삭제:')
    add_steps(doc, [
        '목록에서 "수정" 클릭 → 내용 변경 → "저장"',
        '"삭제" 클릭 → 확인 → 신청 내역도 함께 삭제',
    ])
    add_info_box(doc, '등록된 공지사항은 학생 메뉴(취업박람회/산업체견학/자격증지원 페이지)에 즉시 반영됩니다.')
    add_info_box(doc, '공지사항을 삭제하면 해당 공지의 신청 내역도 함께 삭제됩니다.', 'warning')

    # 코드 관리
    h2(doc, '4.6 코드 관리 (학교 · 전공)')
    body(doc, '회원가입 시 선택 항목으로 표시되는 학교, 전공 정보를 관리합니다.')
    add_screenshot(doc, img('20_admin_codes'), '▲ 코드 관리')
    add_steps(doc, [
        '메뉴 → "코드 관리" 클릭',
        '"학교 추가": 학교명 입력 → 저장 → 회원가입 선택 항목에 자동 추가',
        '목록에서 "수정" / "삭제" 버튼으로 관리',
    ])
    add_info_box(doc, '해당 학교로 등록된 회원이 있는 경우 먼저 회원의 학교 정보를 변경해야 합니다.', 'warning')

    add_page_break(doc)

    # ── 문제 해결 ──────────────────────────────────────────────────
    h1(doc, '부록. 문제 해결 및 문의')

    h2(doc, 'A. 자주 발생하는 문제')
    problems = [
        ('로그인이 안 될 때',
         ['이메일·비밀번호를 정확히 입력했는지 확인',
          'Mac: 이메일 첫 글자 자동 대문자 → 소문자로 직접 입력',
          '브라우저 캐시 삭제 (Ctrl+Shift+R) 후 재시도',
          '다른 브라우저(Chrome, Edge)로 시도']),
        ('화면 데이터가 표시되지 않을 때',
         ['페이지 새로고침 (F5 또는 Ctrl+R)',
          '브라우저 강제 새로고침 (Ctrl+Shift+R)',
          '인터넷 연결 상태 확인']),
        ('회원가입 후 로그인이 안 될 때',
         ['가입 완료 메시지 확인 (오류 메시지가 있었다면 재가입 필요)',
          '가입 시 입력한 이메일·비밀번호로 로그인 시도',
          '관리자에게 계정 확인 요청']),
    ]
    for title_text, steps in problems:
        p = doc.add_paragraph()
        p.add_run(f'■ {title_text}').font.bold = True
        p.runs[0].font.size = Pt(11)
        add_steps(doc, steps)
        doc.add_paragraph()

    h2(doc, 'B. 문의처')
    add_bullets(doc, [
        '이메일: admin@jjobb.kr',
        '사이트: https://jjobb.kr',
        '지원 시간: 평일 09:00 – 18:00',
    ])

    doc.add_paragraph()
    footer = doc.add_paragraph('전북지역 졸업생 네트워크 플랫폼  |  v2.0  |  2026년 2월 26일')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = GRAY
    footer.runs[0].font.size = Pt(9)

    # ── 저장 ──────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f'✅ DOCX 저장 완료: {OUTPUT_FILE}')


if __name__ == '__main__':
    build_manual()

"""
전북지역 졸업생 네트워크 - 개발자 매뉴얼 DOCX 생성
실행: python3 scripts/generate-dev-manual.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_FILE = os.path.join(BASE_DIR, '전북지역_졸업생네트워크_개발자매뉴얼.docx')

BRAND_BLUE  = RGBColor(0x1e, 0x40, 0xaf)
DARK        = RGBColor(0x1f, 0x29, 0x37)
GRAY        = RGBColor(0x6b, 0x72, 0x80)
GREEN       = RGBColor(0x06, 0x60, 0x2f)
RED_DARK    = RGBColor(0x7f, 0x1d, 0x1d)
CODE_BG     = RGBColor(0xf1, 0xf5, 0xf9)


# ─────────────────────────── 헬퍼 ────────────────────────────

def add_page_break(doc):
    doc.add_page_break()


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = BRAND_BLUE
    p.runs[0].font.size = Pt(16)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = DARK
    p.runs[0].font.size = Pt(13)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = DARK
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10.5)
    return p


def code_block(doc, text):
    """코드 블록 스타일의 단락"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # 배경 음영 효과를 위해 shading XML 추가
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F5F9')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    return p


def inline_code(paragraph, text):
    """인라인 코드"""
    run = paragraph.add_run(f' {text} ')
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    return run


def add_bullets(doc, items, indent=0):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
        p.runs[0].font.size = Pt(10.5)


def add_note(doc, text, kind='info'):
    prefix = '💡 ' if kind == 'info' else ('⚠️ ' if kind == 'warning' else '🔴 ')
    color  = RGBColor(0x1e,0x56,0xa0) if kind=='info' else (RGBColor(0x92,0x40,0x0e) if kind=='warning' else RGBColor(0x7f,0x1d,0x1d))
    p = doc.add_paragraph()
    run = p.add_run(f'{prefix}{text}')
    run.font.size = Pt(10)
    run.font.color.rgb = color
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


def add_api_table(doc, rows):
    """
    rows = [(method, endpoint, auth, description), ...]
    """
    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.style = 'Light Shading Accent 1'
    # 헤더
    hdrs = ['Method', 'Endpoint', 'Auth', '설명']
    method_colors = {'GET': RGBColor(0x05,0x78,0x03), 'POST': RGBColor(0x00,0x4e,0xc7),
                     'PUT': RGBColor(0x92,0x40,0x0e), 'DELETE': RGBColor(0x7f,0x1d,0x1d), 'PATCH': RGBColor(0x5b,0x21,0xb6)}
    for i, h in enumerate(hdrs):
        cell = tbl.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    # 데이터
    for ri, (method, endpoint, auth, desc) in enumerate(rows, 1):
        tbl.cell(ri, 0).text = method
        tbl.cell(ri, 1).text = endpoint
        tbl.cell(ri, 2).text = auth
        tbl.cell(ri, 3).text = desc
        # method 색상
        color = method_colors.get(method, GRAY)
        tbl.cell(ri, 0).paragraphs[0].runs[0].font.color.rgb = color
        tbl.cell(ri, 0).paragraphs[0].runs[0].font.bold = True
        for ci in range(4):
            tbl.cell(ri, ci).paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()  # 여백


def add_db_table(doc, table_name, columns):
    """
    columns = [(col_name, type, nullable, description), ...]
    """
    h3(doc, f'▸ {table_name}')
    tbl = doc.add_table(rows=1 + len(columns), cols=4)
    tbl.style = 'Light Shading Accent 1'
    for i, h in enumerate(['컬럼명', '타입', 'Nullable', '설명']):
        cell = tbl.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    for ri, (col, dtype, nullable, desc) in enumerate(columns, 1):
        for ci, val in enumerate([col, dtype, nullable, desc]):
            tbl.cell(ri, ci).text = val
            r = tbl.cell(ri, ci).paragraphs[0].runs[0]
            r.font.size = Pt(9.5)
            if ci == 0:
                r.font.name = 'Courier New'
    doc.add_paragraph()


# ─────────────────────────── 본문 ────────────────────────────

def build_manual():
    doc = Document()

    # 여백
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── 표지 ──────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()

    t1 = doc.add_heading('전북지역 졸업생 네트워크 플랫폼', level=0)
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1.runs[0].font.color.rgb = BRAND_BLUE
    t1.runs[0].font.size = Pt(22)

    t2 = doc.add_heading('개 발 자 매 뉴 얼', level=0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = DARK
    t2.runs[0].font.size = Pt(18)

    doc.add_paragraph()
    info_lines = [
        '버전: v2.0',
        '작성일: 2026년 2월 26일',
        '대상: 백엔드 / 프론트엔드 개발자',
        'GitHub: https://github.com/jsyang9455/graduate-network',
        '라이브: https://jjobb.kr',
    ]
    for line in info_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(10.5)
        p.runs[0].font.color.rgb = GRAY

    add_page_break(doc)

    # ── 목차 ──────────────────────────────────────────────────
    h1(doc, '목   차')
    toc = [
        ('1장', '아키텍처 개요'),
        ('2장', '개발 환경 설정'),
        ('3장', '프로젝트 구조'),
        ('4장', '데이터베이스 스키마'),
        ('5장', 'Backend API 레퍼런스'),
        ('6장', '인증 (JWT)'),
        ('7장', '프론트엔드 구조'),
        ('8장', '배포 (Docker / AWS)'),
        ('9장', '환경 변수'),
        ('10장', 'Nginx 설정'),
        ('부록', '마이그레이션 · 트러블슈팅'),
    ]
    for num, title in toc:
        p = doc.add_paragraph()
        p.add_run(f'{num}  ').font.bold = True
        p.add_run(title).font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    add_page_break(doc)

    # ╔═══════════════════════════════════╗
    # ║ 1장. 아키텍처 개요               ║
    # ╚═══════════════════════════════════╝
    h1(doc, '1장. 아키텍처 개요')

    h2(doc, '1.1 전체 구성')
    add_bullets(doc, [
        '클라이언트: 순수 HTML + CSS + Vanilla JS (SPA 없음, 다중 HTML 파일)',
        '프론트엔드 서버: Nginx (포트 80) — 정적 파일 서빙 + /api/ 역방향 프록시',
        '백엔드 서버: Node.js + Express (포트 5000/5001)',
        '데이터베이스: PostgreSQL 15',
        '컨테이너: Docker Compose (3-tier)',
        '운영 서버: AWS EC2 (Ubuntu 22.04)',
    ])

    h2(doc, '1.2 기술 스택')
    tbl = doc.add_table(rows=7, cols=3)
    tbl.style = 'Light Shading Accent 1'
    for ci, h in enumerate(['레이어', '기술', '버전']):
        tbl.cell(0, ci).text = h
        tbl.cell(0, ci).paragraphs[0].runs[0].font.bold = True
    rows_data = [
        ('Frontend', 'HTML5 / CSS3 / Vanilla JS', '-'),
        ('Backend', 'Node.js + Express', 'v22 / 4.18'),
        ('Database', 'PostgreSQL', '15-alpine'),
        ('Auth', 'JWT (jsonwebtoken)', '9.0'),
        ('Container', 'Docker + Docker Compose', '3.8'),
        ('Reverse Proxy', 'Nginx', '1.25-alpine'),
    ]
    for ri, (layer, tech, ver) in enumerate(rows_data, 1):
        for ci, val in enumerate([layer, tech, ver]):
            tbl.cell(ri, ci).text = val
            tbl.cell(ri, ci).paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    h2(doc, '1.3 요청 흐름')
    body(doc, '브라우저 요청 처리 순서:')
    add_bullets(doc, [
        '브라우저 → Nginx :80',
        '/api/* 요청 → Nginx가 backend:5000 으로 프록시',
        '정적 파일 요청 → Nginx가 직접 서빙',
        'Backend → PostgreSQL 쿼리 실행',
        'JWT 검증은 Express 미들웨어(middleware/auth.js)에서 처리',
    ])

    add_page_break(doc)

    # ╔═══════════════════════════════════╗
    # ║ 2장. 개발 환경 설정              ║
    # ╚═══════════════════════════════════╝
    h1(doc, '2장. 개발 환경 설정')

    h2(doc, '2.1 필수 도구')
    add_bullets(doc, [
        'Node.js v18 이상 (https://nodejs.org)',
        'PostgreSQL 15 (로컬: Postgres.app 또는 Docker)',
        'Git',
        'Docker Desktop (선택, 통합 실행 시)',
    ])

    h2(doc, '2.2 로컬 백엔드 실행')
    code_block(doc, '# 1. 저장소 클론\ngit clone https://github.com/jsyang9455/graduate-network.git\ncd graduate-network/backend')
    code_block(doc, '# 2. 의존성 설치\nnpm install')
    code_block(doc, '# 3. 환경 변수 파일 생성 (backend/.env)\nDB_HOST=localhost\nDB_PORT=5432\nDB_NAME=graduate_network\nDB_USER=<your_pg_user>\nDB_PASSWORD=<your_pg_password>\nJWT_SECRET=your_jwt_secret_key\nJWT_EXPIRE=7d\nPORT=5001\nCORS_ORIGIN=*')
    code_block(doc, '# 4. DB 초기화 (최초 1회)\npsql -U <user> -c "CREATE DATABASE graduate_network;"\npsql -U <user> -d graduate_network -f ../database/schema.sql\npsql -U <user> -d graduate_network -f ../database/seed.sql')
    code_block(doc, '# 5. 서버 실행\nnpm run dev   # nodemon (자동 재시작)\n# 또는\nnode server.js')

    add_note(doc, '로컬 개발 시 백엔드는 포트 5001을 사용합니다. api.js에서 localhost:5001로 자동 감지됩니다.')

    h2(doc, '2.3 로컬 프론트엔드 실행')
    code_block(doc, '# 루트 디렉토리에서\npython3 -m http.server 8080\n# 또는 VS Code Live Server 확장 사용\n# 브라우저: http://localhost:8080')

    h2(doc, '2.4 Docker Compose 전체 실행')
    code_block(doc, 'cd graduate-network\ndocker-compose up -d\n# 확인\ndocker-compose ps')
    add_bullets(doc, [
        'frontend: http://localhost:80',
        'backend:  http://localhost:5000',
        'postgres: localhost:5432',
    ])

    add_page_break(doc)

    # ╔═══════════════════════════════════╗
    # ║ 3장. 프로젝트 구조               ║
    # ╚═══════════════════════════════════╝
    h1(doc, '3장. 프로젝트 구조')

    h2(doc, '3.1 디렉토리 구조')
    code_block(doc,
'''graduate-network/
├── index.html              # 메인(랜딩) 페이지
├── login.html              # 로그인
├── register.html           # 회원가입
├── dashboard.html          # 대시보드 (역할별 분기)
├── jobs.html               # 채용 정보
├── job-fair.html           # 취업박람회
├── industry-visit.html     # 산업체 견학
├── certification-support.html  # 자격증 지원
├── counseling.html         # 진로 상담
├── networking.html         # 동문 네트워킹
├── profile.html            # 내 프로필
├── career.html             # 경력 관리
├── admin-users.html        # 관리자 - 회원 관리
├── admin-jobs.html         # 관리자 - 채용공고 관리
├── admin-board.html        # 관리자 - 게시판 관리
├── admin-announcements.html # 관리자 - 공지사항 관리
├── admin-codes.html        # 관리자 - 코드 관리
├── help.html               # 사용자 도움말
├── css/                    # 스타일시트
│   ├── style.css           # 공통 스타일
│   ├── auth.css            # 인증 페이지
│   ├── dashboard.css       # 대시보드
│   ├── jobs.css            # 채용 정보
│   ├── counseling.css      # 진로 상담
│   └── networking.css      # 네트워킹
├── js/                     # 프론트엔드 스크립트
│   ├── api.js              # API 유틸리티 (base URL, fetch wrapper)
│   ├── auth.js             # 인증 관리 (AuthManager 클래스)
│   ├── main.js             # 메인 페이지
│   ├── login.js            # 로그인 로직
│   ├── register.js         # 회원가입 로직
│   ├── dashboard.js        # 대시보드
│   ├── jobs.js             # 채용 정보
│   ├── counseling.js       # 진로 상담
│   ├── networking.js       # 네트워킹
│   └── admin-users.js      # 관리자 회원관리
├── images/                 # 이미지 에셋
├── backend/                # 백엔드 (Node.js/Express)
│   ├── server.js           # 엔트리포인트
│   ├── package.json
│   ├── Dockerfile
│   ├── config/
│   │   └── database.js     # PostgreSQL 커넥션 풀
│   ├── middleware/
│   │   └── auth.js         # JWT 검증 + 역할 체크
│   ├── routes/             # API 라우터
│   │   ├── auth.js
│   │   ├── users.js
│   │   ├── jobs.js
│   │   ├── announcements.js
│   │   ├── counseling.js
│   │   ├── networking.js
│   │   ├── posts.js
│   │   ├── certificates.js
│   │   └── majors.js
│   └── scripts/
│       ├── migrate.js
│       └── seed.js
├── database/
│   ├── schema.sql          # 테이블 DDL
│   ├── seed.sql            # 초기 데이터
│   └── test-accounts.sql   # 테스트 계정
├── scripts/                # 유틸리티 스크립트
│   ├── capture-screenshots.js
│   ├── generate-manual.py
│   └── generate-dev-manual.py
├── docker-compose.yml
├── nginx.conf
└── Dockerfile              # 프론트엔드 빌드용''')

    add_page_break(doc)

    # ╔═══════════════════════════════════╗
    # ║ 4장. 데이터베이스 스키마          ║
    # ╚═══════════════════════════════════╝
    h1(doc, '4장. 데이터베이스 스키마')

    h2(doc, '4.1 ERD 요약')
    add_bullets(doc, [
        'users ← graduate_profiles (1:1)',
        'users ← company_profiles (1:1)',
        'users → jobs (1:N, company_id)',
        'users ↔ jobs → job_applications (N:M)',
        'users ↔ users → connections (M:M)',
        'users → counseling_sessions (N:1 with counselor)',
        'announcements → announcement_applications (1:N)',
    ])

    h2(doc, '4.2 테이블 상세')

    add_db_table(doc, 'users', [
        ('id',            'SERIAL PK',    'NOT NULL', '사용자 고유 ID'),
        ('email',         'VARCHAR(255)', 'NOT NULL', '로그인 이메일 (UNIQUE)'),
        ('password_hash', 'VARCHAR(255)', 'NOT NULL', 'bcrypt 해시(cost=10)'),
        ('name',          'VARCHAR(100)', 'NOT NULL', '사용자 이름'),
        ('user_type',     'VARCHAR(20)',  'NOT NULL', 'student / graduate / teacher / company / admin'),
        ('phone',         'VARCHAR(20)',  'NULL',     '전화번호'),
        ('school_name',   'VARCHAR(200)', 'NULL',     '학교명'),
        ('is_active',     'BOOLEAN',      'DEFAULT TRUE', '계정 활성화 여부'),
        ('created_at',    'TIMESTAMP',    'DEFAULT NOW()', '가입일'),
        ('last_login',    'TIMESTAMP',    'NULL',     '마지막 로그인'),
    ])

    add_db_table(doc, 'graduate_profiles', [
        ('user_id',           'INT FK→users',  'NOT NULL', '사용자 참조'),
        ('graduation_year',   'INTEGER',        'NOT NULL', '졸업년도'),
        ('major',             'VARCHAR(100)',   'NULL',     '전공'),
        ('current_company',   'VARCHAR(200)',   'NULL',     '재직 회사'),
        ('current_position',  'VARCHAR(100)',   'NULL',     '직위'),
        ('skills',            'TEXT[]',         'NULL',     '스킬 배열'),
        ('is_mentor',         'BOOLEAN',        'DEFAULT FALSE', '멘토 여부'),
    ])

    add_db_table(doc, 'jobs', [
        ('id',            'SERIAL PK',   'NOT NULL', '공고 ID'),
        ('company_id',    'INT FK→users','NOT NULL', '등록자(교사/기업) ID'),
        ('title',         'VARCHAR(255)','NOT NULL', '채용 직무명'),
        ('job_type',      'VARCHAR(50)', 'NULL',     'full-time / part-time / contract / internship'),
        ('status',        'VARCHAR(20)', 'DEFAULT active', 'active / closed / draft'),
        ('deadline',      'DATE',        'NULL',     '지원 마감일'),
        ('views_count',   'INTEGER',     'DEFAULT 0','조회수'),
    ])

    add_db_table(doc, 'announcements', [
        ('id',         'SERIAL PK',   'NOT NULL', '공지 ID'),
        ('type',       'VARCHAR(50)', 'NOT NULL', 'job-fair / industry-visit / certification'),
        ('title',      'VARCHAR(200)','NOT NULL', '제목'),
        ('organizer',  'VARCHAR(100)','NULL',     '주최기관'),
        ('event_date', 'DATE',        'NULL',     '행사일'),
        ('deadline',   'DATE',        'NULL',     '신청 마감일'),
        ('capacity',   'INTEGER',     'NULL',     '모집 인원'),
        ('benefits',   'TEXT[]',      'NULL',     '혜택 배열'),
        ('requirements','TEXT[]',     'NULL',     '준비사항 배열'),
        ('image_url',  'VARCHAR(500)','NULL',     '대표 이미지 URL (선택)'),
        ('detail_url', 'VARCHAR(500)','NULL',     '상세 페이지 URL (선택)'),
        ('is_active',  'BOOLEAN',     'DEFAULT TRUE','활성 여부'),
    ])

    add_db_table(doc, 'announcement_applications', [
        ('id',               'SERIAL PK',           'NOT NULL', '신청 ID'),
        ('announcement_id',  'INT FK→announcements', 'NOT NULL', '공지 참조'),
        ('user_id',          'INT FK→users',         'NULL',     '신청자 (NULL=비로그인)'),
        ('applicant_name',   'VARCHAR(100)',          'NOT NULL', '신청자 이름'),
        ('applicant_phone',  'VARCHAR(20)',           'NOT NULL', '연락처'),
        ('applicant_email',  'VARCHAR(100)',          'NULL',     '이메일'),
        ('status',           'VARCHAR(20)',           'DEFAULT pending', 'pending / approved / rejected'),
    ])

    add_db_table(doc, 'counseling_sessions', [
        ('user_id',     'INT FK→users', 'NOT NULL', '학생 ID'),
        ('counselor_id','INT FK→users', 'NULL',     '교사 ID'),
        ('session_type','VARCHAR(50)',  'NULL',     '진로 / 취업 / 학업 / 기타'),
        ('status',      'VARCHAR(20)',  'DEFAULT scheduled', 'scheduled / completed / cancelled'),
        ('topic',       'TEXT',         'NULL',     '상담 내용'),
        ('notes',       'TEXT',         'NULL',     '답변/교사 메모'),
    ])

    add_page_break(doc)

    # ╔═══════════════════════════════════╗
    # ║ 5장. Backend API 레퍼런스        ║
    # ╚═══════════════════════════════════╝
    h1(doc, '5장. Backend API 레퍼런스')

    h2(doc, '5.0 공통 사항')
    body(doc, 'Base URL: http://localhost:5001/api  (로컬) | https://jjobb.kr/api  (운영)')
    add_bullets(doc, [
        'Content-Type: application/json',
        'Authorization: Bearer <JWT_TOKEN>  (인증 필요 엔드포인트)',
        '성공: HTTP 200/201, 실패: 400 / 401 / 403 / 404 / 500',
    ])
    code_block(doc, '# 헬스 체크\nGET /api/health\n→ {"status":"OK","message":"...","timestamp":"..."}')

    h2(doc, '5.1 인증 (POST /api/auth)')
    add_api_table(doc, [
        ('POST',   '/api/auth/register',        '없음',      '회원가입'),
        ('POST',   '/api/auth/login',           '없음',      '로그인 → JWT 반환'),
        ('GET',    '/api/auth/me',              'JWT 필요',  '현재 유저 정보 조회'),
        ('POST',   '/api/auth/change-password', 'JWT 필요',  '비밀번호 변경'),
    ])
    h3(doc, 'POST /api/auth/register  요청 바디')
    code_block(doc, '{\n  "email": "user@example.com",\n  "password": "min6chars",\n  "name": "홍길동",\n  "user_type": "student",  // student|graduate|teacher|company|admin\n  "phone": "010-1234-5678",  // 선택\n  "school_name": "○○고등학교"  // 선택\n}')
    h3(doc, 'POST /api/auth/login  요청/응답')
    code_block(doc, '// 요청\n{"email":"user@example.com","password":"1234"}\n// 응답\n{"message":"Login successful","token":"eyJ...","user":{"id":1,"email":"...","name":"...","user_type":"..."}}')

    h2(doc, '5.2 사용자 (GET /api/users)')
    add_api_table(doc, [
        ('GET',    '/api/users',              'admin',     '전체 회원 목록'),
        ('GET',    '/api/users/:id',          'JWT 필요',  '특정 사용자 조회'),
        ('PUT',    '/api/users/profile',      'JWT 필요',  '내 프로필 수정'),
        ('PUT',    '/api/users/:id',          'admin',     '관리자 - 회원 수정'),
        ('DELETE', '/api/users/:id',          'admin',     '관리자 - 회원 삭제'),
        ('GET',    '/api/users/graduate-profile/:userId', 'JWT 필요', '졸업생 프로필 조회'),
        ('PUT',    '/api/users/graduate-profile', 'JWT 필요', '졸업생 프로필 수정'),
    ])

    h2(doc, '5.3 채용공고 (GET /api/jobs)')
    add_api_table(doc, [
        ('GET',    '/api/jobs',                '없음',          '채용공고 목록 (필터: status, location 등)'),
        ('GET',    '/api/jobs/:id',            '없음',          '채용공고 상세'),
        ('POST',   '/api/jobs',                'teacher/admin', '채용공고 등록'),
        ('PUT',    '/api/jobs/:id',            'teacher/admin', '채용공고 수정'),
        ('DELETE', '/api/jobs/:id',            'admin',         '채용공고 삭제'),
        ('POST',   '/api/jobs/:id/apply',      'JWT 필요',      '채용공고 지원'),
        ('GET',    '/api/jobs/my/applications','JWT 필요',      '내 지원 내역'),
    ])

    h2(doc, '5.4 공지사항 (GET /api/announcements)')
    add_api_table(doc, [
        ('GET',    '/api/announcements/:type',           '없음',    '유형별 공지 목록 (job-fair|industry-visit|certification)'),
        ('GET',    '/api/announcements/detail/:id',      '없음',    '공지 상세'),
        ('POST',   '/api/announcements',                 'admin',   '공지 등록'),
        ('PUT',    '/api/announcements/:id',             'admin',   '공지 수정'),
        ('DELETE', '/api/announcements/:id',             'admin',   '공지 삭제'),
        ('POST',   '/api/announcements/apply',           'JWT 권장','참가 신청'),
        ('GET',    '/api/announcements/my-applications', 'JWT 필요','내 신청 내역'),
        ('GET',    '/api/announcements/applications/all','admin',   '전체 신청 목록'),
        ('PUT',    '/api/announcements/applications/:id/status','admin','신청 상태 변경'),
    ])

    h2(doc, '5.5 진로 상담 (GET /api/counseling)')
    add_api_table(doc, [
        ('GET',    '/api/counseling',    'JWT 필요',      '상담 목록 (역할별 필터)'),
        ('POST',   '/api/counseling',    'JWT 필요',      '상담 신청'),
        ('PUT',    '/api/counseling/:id','teacher/admin', '상담 승인/거절/답변'),
        ('DELETE', '/api/counseling/:id','JWT 필요',      '상담 삭제'),
    ])

    h2(doc, '5.6 네트워킹 (GET /api/networking)')
    add_api_table(doc, [
        ('GET',  '/api/networking/connections',     'JWT 필요', '연결 목록'),
        ('GET',  '/api/networking/requests',        'JWT 필요', '받은 요청 목록'),
        ('POST', '/api/networking/connect/:userId', 'JWT 필요', '네트워크 연결 요청'),
        ('PUT',  '/api/networking/requests/:id',    'JWT 필요', '요청 수락/거절'),
        ('GET',  '/api/networking/mentors',         'JWT 필요', '멘토 목록'),
    ])

    h2(doc, '5.7 게시판/코드 관리')
    add_api_table(doc, [
        ('GET',    '/api/posts',          '없음',  '게시글 목록'),
        ('POST',   '/api/posts',          'JWT',   '게시글 작성'),
        ('PUT',    '/api/posts/:id',      'JWT',   '게시글 수정'),
        ('DELETE', '/api/posts/:id',      'JWT',   '게시글 삭제'),
        ('GET',    '/api/majors',         '없음',  '전공 목록'),
        ('POST',   '/api/majors',         'admin', '전공 추가'),
        ('PUT',    '/api/majors/:id',     'admin', '전공 수정'),
        ('DELETE', '/api/majors/:id',     'admin', '전공 삭제'),
    ])

    add_page_break(doc)

    # ╔═══════════════════════════════════╗
    # ║ 6장. 인증 (JWT)                  ║
    # ╚═══════════════════════════════════╝
    h1(doc, '6장. 인증 (JWT)')

    h2(doc, '6.1 토큰 생성')
    code_block(doc, '// backend/routes/auth.js\njwt.sign(\n  { id: user.id, email: user.email, user_type: user.user_type },\n  process.env.JWT_SECRET,\n  { expiresIn: process.env.JWT_EXPIRE || "7d" }\n);')

    h2(doc, '6.2 토큰 검증 미들웨어')
    code_block(doc, '// backend/middleware/auth.js\nconst auth = async (req, res, next) => {\n  const token = req.header("Authorization")?.replace("Bearer ", "");\n  if (!token) return res.status(401).json({ error: "Authentication required" });\n  const decoded = jwt.verify(token, process.env.JWT_SECRET);\n  req.user = decoded;  // { id, email, user_type }\n  next();\n};')

    h2(doc, '6.3 역할 기반 접근 제어')
    code_block(doc, '// 사용 예시 (routes/*.js)\nrouter.post("/", auth, checkRole("teacher", "admin"), handler);\n\n// checkRole 구현\nconst checkRole = (...roles) => (req, res, next) => {\n  if (!roles.includes(req.user.user_type))\n    return res.status(403).json({ error: "Access denied" });\n  next();\n};')

    add_bullets(doc, [
        'student: 본인 데이터 읽기/쓰기, 지원, 상담 신청',
        'graduate: student와 동일 + 졸업생 프로필',
        'teacher: 채용공고 등록, 상담 답변',
        'admin: 모든 API 접근 가능',
    ])

    h2(doc, '6.4 프론트엔드 토큰 관리')
    code_block(doc, '// js/api.js\nlocalStorage.setItem("token", token);\nlocalStorage.setItem("graduateNetwork_user", JSON.stringify(user));\n\n// API 호출 시 자동 첨부\nfetch(url, {\n  headers: { Authorization: `Bearer ${localStorage.getItem("token")}` }\n});')

    add_page_break(doc)

    # ╔═══════════════════════════════════╗
    # ║ 7장. 프론트엔드 구조             ║
    # ╚═══════════════════════════════════╝
    h1(doc, '7장. 프론트엔드 구조')

    h2(doc, '7.1 API 유틸리티 (js/api.js)')
    body(doc, 'API_BASE_URL은 hostname을 감지하여 자동 설정됩니다:')
    code_block(doc, '// js/api.js\nconst API_BASE_URL = \n  window.location.hostname === "localhost" || window.location.hostname === "127.0.0.1"\n  ? "http://localhost:5001/api"\n  : "/api";  // 운영: Nginx 프록시 경유')
    add_note(doc, 'HTML 파일 내에 API URL을 하드코딩하지 마세요. 반드시 api.js의 API_BASE_URL 변수를 사용하세요.', 'warning')

    h2(doc, '7.2 인증 클래스 (js/auth.js)')
    body(doc, 'AuthManager 클래스가 로그인 상태 관리를 담당합니다:')
    add_bullets(doc, [
        'isLoggedIn(): localStorage의 token 여부로 판단',
        'getCurrentUser(): graduateNetwork_user 파싱 반환',
        'login(userData, token): localStorage에 저장',
        'logout(): localStorage 초기화 + 로그인 페이지 이동',
        'checkAuth(): 페이지 로드 시 JWT 유효성 검증',
    ])

    h2(doc, '7.3 페이지별 역할 접근 제어')
    code_block(doc, '// 각 페이지 JS에서 역할 체크 예시\nconst user = auth.getCurrentUser();\nif (!user || user.user_type !== "admin") {\n  window.location.href = "/login.html";\n  return;\n}')

    h2(doc, '7.4 신규 HTML 페이지 추가 시 체크리스트')
    add_bullets(doc, [
        '헤더에 api.js, auth.js 스크립트 포함',
        'API 호출 시 API_BASE_URL 변수 사용 (localhost:5001 하드코딩 금지)',
        '역할 체크 로직 추가 (관리자 전용 페이지 등)',
        'Nginx 컨테이너 재빌드 필요: docker-compose build --no-cache frontend',
    ])

    add_page_break(doc)

    # ╔═══════════════════════════════════╗
    # ║ 8장. 배포 (Docker / AWS)         ║
    # ╚═══════════════════════════════════╝
    h1(doc, '8장. 배포 (Docker / AWS)')

    h2(doc, '8.1 Docker Compose 서비스 구성')
    tbl = doc.add_table(rows=4, cols=5)
    tbl.style = 'Light Shading Accent 1'
    for ci, h in enumerate(['서비스', '컨테이너명', '이미지', '포트', '역할']):
        tbl.cell(0, ci).text = h
        tbl.cell(0, ci).paragraphs[0].runs[0].font.bold = True
    svc_data = [
        ('postgres',  'graduate-network-db',       'postgres:15-alpine', '5432:5432', 'PostgreSQL DB'),
        ('backend',   'graduate-network-backend',  'custom(Dockerfile)',  '5000:5000', 'Express API'),
        ('frontend',  'graduate-network-frontend', 'nginx:1.25-alpine',  '80:80',     'Nginx + 정적파일'),
    ]
    for ri, row in enumerate(svc_data, 1):
        for ci, val in enumerate(row):
            tbl.cell(ri, ci).text = val
            tbl.cell(ri, ci).paragraphs[0].runs[0].font.size = Pt(9.5)
    doc.add_paragraph()

    h2(doc, '8.2 AWS EC2 최초 배포')
    code_block(doc, '# SSH 접속\nssh -i your-key.pem ubuntu@<EC2_IP>')
    code_block(doc, '# Docker 설치 (Ubuntu 22.04)\nsudo apt update && sudo apt install -y docker.io docker-compose\nsudo usermod -aG docker ubuntu && newgrp docker')
    code_block(doc, '# 저장소 클론\ncd /home/ubuntu\ngit clone https://github.com/jsyang9455/graduate-network.git\ncd graduate-network')
    code_block(doc, '# 빌드 및 실행\ndocker-compose build --no-cache\ndocker-compose up -d\n# 확인\ndocker-compose ps\ndocker-compose logs -f backend')

    h2(doc, '8.3 코드 변경 후 재배포')
    add_api_table(doc, [
        ('변경 유형',            '명령어',                                                        '', ''),
    ])
    # 별도 표로 처리
    deploy_tbl = doc.add_table(rows=5, cols=2)
    deploy_tbl.style = 'Light Shading Accent 1'
    deploy_tbl.cell(0, 0).text = '변경 유형'
    deploy_tbl.cell(0, 1).text = '명령어'
    deploy_tbl.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    deploy_tbl.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    deploy_rows = [
        ('HTML/CSS/JS 변경',    'git pull && docker-compose build --no-cache frontend && docker-compose up -d frontend'),
        ('백엔드 JS 변경',      'git pull && docker-compose build --no-cache backend && docker-compose up -d backend'),
        ('DB 스키마 변경',      'git pull && docker-compose exec postgres psql -U postgres graduate_network -f /migration.sql'),
        ('전체 재빌드',         'git pull && docker-compose build --no-cache && docker-compose up -d'),
    ]
    for ri, (change, cmd) in enumerate(deploy_rows, 1):
        deploy_tbl.cell(ri, 0).text = change
        deploy_tbl.cell(ri, 1).text = cmd
        for ci in range(2):
            deploy_tbl.cell(ri, ci).paragraphs[0].runs[0].font.size = Pt(9.5)
    doc.add_paragraph()

    add_note(doc, 'HTML/CSS/JS 파일 변경 후 git pull만으로는 반영되지 않습니다. 반드시 docker-compose build --no-cache frontend 실행 후 up -d 해야 합니다.', 'warning')

    h2(doc, '8.4 DB 마이그레이션')
    code_block(doc, '# AWS에서 마이그레이션 SQL 실행\ndocker-compose exec postgres psql -U postgres -d graduate_network \\\n  -c "ALTER TABLE announcements ADD COLUMN IF NOT EXISTS image_url VARCHAR(500);"')
    add_note(doc, '기존 데이터 보존을 위해 항상 IF NOT EXISTS / IF NOT EXISTS 구문 사용. DROP 절대 금지.')

    h2(doc, '8.5 컨테이너 관리 명령어')
    code_block(doc, '# 로그 확인\ndocker-compose logs -f backend\ndocker-compose logs -f frontend\n\n# 컨테이너 접속\ndocker-compose exec backend sh\ndocker-compose exec postgres psql -U postgres -d graduate_network\n\n# 재시작\ndocker-compose restart backend\n\n# 전체 중지\ndocker-compose down\n\n# 볼륨 포함 초기화 (DB 데이터 삭제)\ndocker-compose down -v')

    add_page_break(doc)

    # ╔═══════════════════════════════════╗
    # ║ 9장. 환경 변수                   ║
    # ╚═══════════════════════════════════╝
    h1(doc, '9장. 환경 변수')

    h2(doc, '9.1 backend/.env (로컬 개발)')
    code_block(doc, 'NODE_ENV=development\nPORT=5001\n\n# PostgreSQL\nDB_HOST=localhost\nDB_PORT=5432\nDB_NAME=graduate_network\nDB_USER=<pg_username>\nDB_PASSWORD=<pg_password>\n\n# JWT\nJWT_SECRET=<임의의_긴_문자열_32자_이상>\nJWT_EXPIRE=7d\n\n# CORS\nCORS_ORIGIN=*')

    h2(doc, '9.2 docker-compose.yml (운영)')
    code_block(doc, '# backend 서비스 environment\nNODE_ENV: production\nPORT: 5000\nDB_HOST: postgres           # 컨테이너 서비스명\nDB_PORT: 5432\nDB_NAME: graduate_network\nDB_USER: postgres\nDB_PASSWORD: postgres\nJWT_SECRET: graduate_network_secret_key_2026_jeonju_tech\nJWT_EXPIRE: 7d\nCORS_ORIGIN: "*"')

    add_note(doc, '운영 환경에서는 JWT_SECRET을 충분히 복잡한 임의 값으로 변경하고, DB_PASSWORD도 강력한 값으로 변경하세요.', 'warning')

    add_page_break(doc)

    # ╔═══════════════════════════════════╗
    # ║ 10장. Nginx 설정                 ║
    # ╚═══════════════════════════════════╝
    h1(doc, '10장. Nginx 설정')

    h2(doc, '10.1 nginx.conf 핵심 설정')
    code_block(doc, '# /api/ 요청 → 백엔드 프록시\nlocation /api/ {\n  proxy_pass http://backend:5000/api/;\n  proxy_set_header Host $host;\n  proxy_set_header X-Real-IP $remote_addr;\n}\n\n# SPA 폴백 없음 (다중 HTML 파일 구조)\nlocation / {\n  try_files $uri $uri/ =404;\n}\n\n# JS/CSS 캐시 비활성화\nlocation ~* \\.(js|css)$ {\n  expires -1;\n  add_header Cache-Control "no-store, no-cache";\n}\n\n# 이미지 1년 캐시\nlocation ~* \\.(jpg|jpeg|png|gif|ico)$ {\n  expires 1y;\n  add_header Cache-Control "public, immutable";\n}')

    add_note(doc, 'Docker 내부 네트워크에서 backend는 컨테이너명(backend)으로 참조됩니다. proxy_pass에 localhost:5000 사용 불가.')

    add_page_break(doc)

    # ╔═══════════════════════════════════╗
    # ║ 부록. 마이그레이션/트러블슈팅    ║
    # ╚═══════════════════════════════════╝
    h1(doc, '부록. 마이그레이션 및 트러블슈팅')

    h2(doc, 'A. 안전한 마이그레이션 패턴')
    code_block(doc, '-- 새 테이블 추가 (안전)\nCREATE TABLE IF NOT EXISTS new_table (...);\n\n-- 컬럼 추가 (안전)\nDO $$ BEGIN\n  IF NOT EXISTS (\n    SELECT 1 FROM information_schema.columns\n    WHERE table_name=\'table\' AND column_name=\'col\'\n  ) THEN\n    ALTER TABLE table ADD COLUMN col VARCHAR(100);\n  END IF;\nEND $$;\n\n-- 절대 사용 금지 (데이터 손실)\nDROP TABLE table;\nDROP COLUMN col;')

    h2(doc, 'B. 자주 발생하는 개발 오류')

    issues = [
        ('HTML 변경이 운영 서버에 반영 안 됨',
         ['git pull 만으로는 부족합니다.',
          'docker-compose build --no-cache frontend 실행',
          'docker-compose up -d frontend 실행']),
        ('CORS 오류 (Access-Control-Allow-Origin)',
         ['HTML 파일에서 localhost:5001 하드코딩 확인',
          'api.js의 API_BASE_URL 변수를 반드시 사용할 것',
          'docker-compose.yml의 CORS_ORIGIN 값 확인']),
        ('401 Unauthorized',
         ['JWT 토큰 만료 여부 확인 (기본 7일)',
          'localStorage의 token 값 확인 (개발자도구 → Application)',
          'JWT_SECRET이 .env와 docker-compose.yml에서 일치하는지 확인']),
        ('DB 컬럼 누락 오류 (column does not exist)',
         ['AWS DB에 마이그레이션 SQL이 실행되지 않은 상태',
          'docker-compose exec postgres psql 로 수동 ALTER TABLE 실행',
          '로컬 DB와 AWS DB 스키마 불일치 점검']),
        ('Docker 빌드 후에도 구버전 파일',
         ['docker system prune 으로 캐시 제거',
          'docker-compose build --no-cache 옵션 확인']),
    ]
    for title_text, steps in issues:
        p = doc.add_paragraph()
        p.add_run(f'■ {title_text}').font.bold = True
        p.runs[0].font.size = Pt(11)
        add_bullets(doc, steps)
        doc.add_paragraph()

    h2(doc, 'C. 유용한 진단 명령어')
    code_block(doc, '# 서버 상태 확인\ncurl http://localhost:5001/api/health\n\n# DB 직접 접속 (로컬)\nexport PATH="/Applications/Postgres.app/Contents/Versions/latest/bin:$PATH"\npsql -U <user> -d graduate_network\n\n# 테이블 목록\n\\dt\n\n# 컬럼 확인\n\\d users\n\n# 사용자 확인\nSELECT id, email, user_type, is_active FROM users LIMIT 10;\n\n# Docker 로그\ndocker-compose logs -f --tail=100 backend\n\n# 컨테이너 상태\ndocker-compose ps')

    h2(doc, 'D. 테스트 계정 (로컬 개발)')
    tbl2 = doc.add_table(rows=4, cols=3)
    tbl2.style = 'Light Shading Accent 1'
    for ci, h in enumerate(['역할', '이메일', '비밀번호']):
        tbl2.cell(0, ci).text = h
        tbl2.cell(0, ci).paragraphs[0].runs[0].font.bold = True
    test_accounts = [
        ('관리자 (admin)',  'admin@jeonjutech.edu',   'manual123'),
        ('교사 (teacher)', 'teacher.kim@example.com', 'manual123'),
        ('학생 (student)', 'jung.yuna@example.com',  'manual123'),
    ]
    for ri, (role, email, pw) in enumerate(test_accounts, 1):
        for ci, val in enumerate([role, email, pw]):
            tbl2.cell(ri, ci).text = val
            tbl2.cell(ri, ci).paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    add_note(doc, '위 계정들은 로컬 개발/테스트 전용입니다. 운영 서버(jjobb.kr)에서는 실제 비밀번호로 변경하세요.', 'warning')

    # 하단
    doc.add_paragraph()
    footer = doc.add_paragraph('전북지역 졸업생 네트워크 플랫폼  |  개발자 매뉴얼 v2.0  |  2026년 2월 26일')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = GRAY
    footer.runs[0].font.size = Pt(9)

    doc.save(OUTPUT_FILE)
    print(f'✅ 개발자 매뉴얼 저장 완료: {OUTPUT_FILE}')


if __name__ == '__main__':
    build_manual()
